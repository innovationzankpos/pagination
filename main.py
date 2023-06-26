import math
import os
import tempfile
from escpos.printer import Serial
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date, datetime
from itertools import count


import PyQt5
import pyodbc
import win32ui
from PyQt5 import QtWidgets, QtCore, QtGui
from PyQt5.QtCore import QEvent, Qt, QObject, QSortFilterProxyModel
from PyQt5.QtGui import QStandardItem, QStandardItemModel, QFont, QCursor, QTextCharFormat, QColor, QTextCursor, \
    QTextDocument
from PyQt5.QtWidgets import QMainWindow, QHeaderView, QAbstractItemView, QTableView, QLineEdit, QStyledItemDelegate, \
    QStyleOptionViewItem, QTextEdit, QMessageBox

from pagination_ui import Ui_MainWindow


# Establish a connection
conn = pyodbc.connect(
    "Driver={SQL Server};"
    "Server={localhost\sqlexpress2014};"
    "Database={POWERPOS};"
)

cursor = conn.cursor()

cursor.execute("SELECT itemcode as ITEMCODE, itemname as ITEMNAME, department as DEPARTMENT, uom as UOM, sellingprice as PRICE, sp_selling as WHOLESALE, end_qty as BAL FROM items ORDER BY itemname ASC")

result = cursor.fetchall()

rows_per_page = 300
current_page = 0
total_pages = 0
isUsed = False
search_result = []
print_filename = ""


class MyEventFilter(QObject):
    def __init__(self, window):
        super().__init__()
        self.window = window

    def eventFilter(self, watched, event):
        if event.type() == QEvent.KeyPress and (event.key() == Qt.Key_Return or event.key() == Qt.Key_Enter):
            print("Triggered!")
            if isinstance(watched, QLineEdit):
                print("Enter key pressed in line edit!")
                self.window.search()
            elif isinstance(watched, QTableView):
                print("Enter key pressed in table view!")
                self.window.enter()
        # Continue processing event as intended
        return super().eventFilter(watched, event)


class MainWindow(QMainWindow, Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)

        self.setFixedSize(1040, 630)
        # self.setMouseTracking(True)  # Enable mouse tracking
        # self.setCursor(Qt.ArrowCursor)  # Set the cursor to arrow

        self.searchTxt.setStyleSheet("""
            QLineEdit {
                border: 1px solid gray; 
                font-size: 24px; 
                font-weight: bold;
                padding: 5px;
            }
            QLineEdit::focus,
            QLineEdit::hover {
                border: 3px solid black;
            }
        """)
        # Connect the textChanged signal to the custom method
        self.searchTxt.textChanged.connect(self.search_isEmpty)
        self.searchBtn.clicked.connect(self.search)
        self.searchBtn.setStyleSheet("""
            QPushButton{
                background-color: #880808;
                color: white;
            }
            QPushButton:hover {
                background-color: red;
                border: 2px solid white;
                padding: 4px;
            }
        """)
        self.itemBtn.clicked.connect(self.item)
        self.itemBtn.setStyleSheet("""
            QPushButton {
                background-color: black;
                color: white;
            }
            QPushButton:hover {
                background-color: #3B3B3B;
                border: 2px solid white;
                padding: 4px;
            }
        """)
        self.brandBtn.clicked.connect(self.brand)
        self.brandBtn.setStyleSheet("""
            QPushButton#brandBtn {
                background-color : black;
                color: white;
            }
            QPushButton#brandBtn:hover {
                background-color: #3B3B3B;
                border: 2px solid white;
                padding: 4px;
            }
        """)
        self.printBtn.clicked.connect(self.print)
        self.printBtn.setStyleSheet("""
            QPushButton#printBtn {
                background-color : #524532;
                color: white;
            }
            QPushButton#printBtn:hover {
                background-color: #6E5C43;
                border: 2px solid white;
                padding: 4px;
            }
        """)
        self.prevBtn.clicked.connect(self.prev)
        self.prevBtn.setStyleSheet("""
            QPushButton{
                background-color: #5D8AA8;
                color: darkgray;
            }
            QPushButton::hover{
                background-color: #4169E1;
                border: 2px solid white;
                padding: 4px;
            }
        """)
        self.prevBtn.setDisabled(True)
        self.nextBtn.clicked.connect(self.next)
        self.nextBtn.setStyleSheet("""
            QPushButton{
                background-color: blue;
                color: white;
            }
            QPushButton::hover{
                background-color: #4169E1;
                border: 2px solid white;
                padding: 4px;
            }
        """)
        self.totalLabel.setStyleSheet("font-size: 14px; font-weight: bold; color: black;")
        self.enterBtn.clicked.connect(self.enter)
        self.enterBtn.setStyleSheet("""
            QPushButton {
                background-color : green;
                color: white;
            }
            QPushButton:hover {
                background-color: #00D100;
                color: black;
                border: 2px solid white;
                padding: 4px;
            }
        """)
        self.closeBtn.clicked.connect(self.closeUI)
        self.closeBtn.setStyleSheet("""
            QPushButton {
                background-color: #880808;
                color: white;
            }
            QPushButton:hover {
                background-color: red;
                border: 2px solid white;
                padding: 4px;
            }
        """)

        self.model = QStandardItemModel()

        # Set the font size
        self.font = QFont()
        # font.setBold(True)
        self.font.setPixelSize(18)  # Replace 12 with the desired font size
        self.font.setWeight(60)

        # Create a QTableView and set its model
        self.tableView.setModel(self.model)

        # Install the event filter
        self.event_filter = MyEventFilter(self)
        self.searchTxt.installEventFilter(self.event_filter)
        self.tableView.installEventFilter(self.event_filter)

        # Hide the vertical header
        self.tableView.verticalHeader().setVisible(False)

        # Set the selection mode to select rows
        self.tableView.setSelectionMode(QAbstractItemView.SingleSelection)
        self.tableView.setSelectionBehavior(QAbstractItemView.SelectRows)

        # Set the focus policy to TabFocus for the table view
        self.tableView.setFocusPolicy(Qt.StrongFocus)

        # Set the table view as non-editable
        self.tableView.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)

        # # Hide the horizontal scroll bar
        self.tableView.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        # Populate the table with data
        self.populate_tableview(current_page)

        # Select the first row
        self.tableView.selectRow(0)

        # Set focus to tableView with a QTimer to ensure the widget is ready
        QtCore.QTimer.singleShot(0, self.tableView.setFocus)

        # Connect the focusChanged signal to a slot
        QtWidgets.QApplication.instance().focusChanged.connect(self.on_focusChanged)

    def search(self):

        print("Searched!")

        # Retrieve the text from the QLineEdit
        input_text = self.searchTxt.text()
        print("Text entered:", input_text)

        if input_text != "":
            global search_result
            global isUsed

            isUsed = True

            # Execute a query to check if the text is a substring of the column in the database
            cursor.execute(
                "SELECT itemcode as ITEMCODE, itemname as ITEMNAME, department as DEPARTMENT, uom as UOM, sellingprice as PRICE, sp_selling as WHOLESALE, end_qty as BAL FROM items WHERE itemname LIKE ? ORDER BY itemname ASC",
                ('%' + input_text + '%',))
            search_result = cursor.fetchall()

            if len(search_result) > 0:
                self.prevBtn.setDisabled(True)
                self.prevBtn.setStyleSheet("background-color: #5D8AA8; color: darkgray;")
                self.nextBtn.setDisabled(False)
                self.nextBtn.setStyleSheet("background-color: blue; color: white;")
                print("Text is a substring in the database")
                global current_page

                current_page = 0
                self.model.clear()
                self.populate_tableview(current_page)
                self.searchTxt.clearFocus()
            else:
                print("Text is not a substring in the database")

        # cursor.close()

    def item(self):

        # Retrieve the text from the QLineEdit
        input_text = self.searchTxt.text()
        print("Text entered:", input_text)

        if input_text != "":
            global search_result
            global isUsed

            isUsed = True

            # Execute a query to check if the text is a substring of the column in the database
            cursor.execute(
                "SELECT itemcode as ITEMCODE, itemname as ITEMNAME, department as DEPARTMENT, uom as UOM, sellingprice as PRICE, sp_selling as WHOLESALE, end_qty as BAL FROM items WHERE itemname = ?",
                input_text)
            search_result = cursor.fetchall()

            if len(search_result) > 0:
                self.prevBtn.setDisabled(True)
                self.prevBtn.setStyleSheet("background-color: #5D8AA8; color: darkgray;")
                self.nextBtn.setDisabled(False)
                self.nextBtn.setStyleSheet("background-color: blue; color: white;")
                print("Text has match/es in the database")
                global current_page

                current_page = 0
                self.model.clear()
                self.populate_tableview(current_page)
            else:
                print("Text doesn't have match in the database")

    def brand(self):

        # Retrieve the text from the QLineEdit
        input_text = self.searchTxt.text()
        print("Text entered:", input_text)

        if input_text != "":
            global search_result
            global isUsed

            isUsed = True

            # Execute a query to check if the text is a substring of the column in the database
            cursor.execute(
                "SELECT itemcode as ITEMCODE, itemname as ITEMNAME, department as DEPARTMENT, uom as UOM, sellingprice as PRICE, sp_selling as WHOLESALE, end_qty as BAL FROM items WHERE brand = ?",
                input_text)
            search_result = cursor.fetchall()

            if len(search_result) > 0:
                self.prevBtn.setDisabled(True)
                self.prevBtn.setStyleSheet("background-color: #5D8AA8; color: darkgray;")
                self.nextBtn.setDisabled(False)
                self.nextBtn.setStyleSheet("background-color: blue; color: white;")
                print("Text is a match in the database")
                global current_page

                current_page = 0
                self.model.clear()
                self.populate_tableview(current_page)
            else:
                print("Text is not a match in the database")

    def print(self):

        dlg = QMessageBox(self)
        dlg.setWindowTitle("Confirmation Message")
        dlg.setText("Are you sure you want to print it?")
        dlg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        dlg.setIcon(QMessageBox.Question)
        button = dlg.exec()

        if button == QMessageBox.Yes:
            print("Yes!")

            # Get the selection model from the table view
            selection_model = self.tableView.selectionModel()

            if selection_model.hasSelection():

                # Get the selected rows
                selected_rows = selection_model.selectedRows()

                # Print the values of all cells in the selected rows
                for index in selected_rows:
                    row = index.row()
                    items = []
                    for column in range(self.model.columnCount()):
                        # Index of the cell in the model
                        cell_index = self.model.index(row, column)

                        # Value of the cell
                        cell_value = self.model.data(cell_index)
                        # print(cell_value)
                        items.append(str(cell_value).strip())

                print(items)
                # self.format_print()
                doc = self.format_print(items)

                # Save the document
                doc.save("Printed/" + print_filename)
                self.closeUI()
            else:
                dlg = QMessageBox(self)
                dlg.setWindowTitle("Warning Message")
                dlg.setText("Please select a row!")
                dlg.setStandardButtons(QMessageBox.Ok)
                dlg.setIcon(QMessageBox.Question)
                button = dlg.exec()

                if button == QMessageBox.Ok:
                    print("Ok!")
        else:
            print("No!")

    def format_print(self, selected):

        keys = ['ITEMCODE:', 'ITEMNAME:', 'DEPARTMENT:', 'UOM:', 'PRICE:', 'W.SALE PRICE:', 'BAL:']

        # Create a new document
        document = Document()

        # Set font size for the entire document
        document.styles['Normal'].font.size = Pt(10)

        # Set margins for the document (in inches)
        sections = document.sections
        for section in sections:
            section.left_margin = Inches(0.2)
            section.right_margin = Inches(0.2)
            section.top_margin = Inches(0.2)
            section.bottom_margin = Inches(0.2)

        # Set page size for the document
        section = document.sections[0]
        section.page_width = Inches(1.89)
        section.page_height = Inches(8.26)

        # Add shop name
        shop_name = 'ZANK POS ENTERPRISES'
        document.add_paragraph(shop_name)

        document.add_paragraph("--------------------------------")

        current_year = date.today().year
        print(current_year)

        print_id = 0
        print_id_txt = ""

        # Read the config file
        with open("config.txt", 'r') as file:
            config_lines = file.readlines()

        # Find the line with the setting key and extract the value
        for line in config_lines:
            line = line.strip()
            if line.startswith('incrementing_ID'):
                _, value = line.split(':')
                print_id = value.strip()
                print_id_txt = str(current_year) + "000" + str(print_id)
                document.add_paragraph(f"Print ID: {print_id_txt}")
                document.add_paragraph("\n")
                print("Incrementing_ID: " + print_id)

        # Find the line with the setting key and update the value
        for i in range(len(config_lines)):
            line = config_lines[i].strip()
            if line.startswith('incrementing_ID'):
                config_lines[i] = f"{'incrementing_ID'}: {int(print_id) + 1}\n"
                break

        # Write the updated content back to the file
        with open("config.txt", 'w') as file:
            file.writelines(config_lines)

        # Add item details
        item_details = {k: v for k, v in zip(keys, selected)}

        for key, value in item_details.items():
            document.add_paragraph(key)

            if any(word in key for word in ["PRICE", "W.SALE PRICE", "BAL"]):
                document.add_paragraph("Php " + value)
            else:
                document.add_paragraph(value)

        document.add_paragraph("\n")

        current_date = date.today()
        formatted_date = current_date.strftime("%m/%d/%Y")

        # Add closing statement
        parag1 = document.add_paragraph('Thank you!')

        # Set font style for paragraph 1
        for run in parag1.runs:
            run.italic = True

        document.add_paragraph("--------------------------------")
        parag2 = document.add_paragraph(f"Date: {formatted_date}")

        # Set font style for paragraph 1
        for run in parag2.runs:
            run.italic = True

        # Set alignment for specific paragraphs
        for paragraph in document.paragraphs:
            if any(word in paragraph.text for word in ["ZANK", "Print", "Thank", "Date:"]):
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif any(word in paragraph.text for word in ["ITEM", "DEPARTMENT", "UOM", "PRICE", "BAL"]):
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        current_datetime = datetime.now()
        formatted_datetime = current_datetime.strftime("%Y-%m-%d-%H-%M-%S-%f")
        print("DateTime: " + formatted_datetime)
        global print_filename
        print_filename = formatted_datetime + "ID" + print_id_txt + ".docx"
        # print_filename = str(current_datetime).replace(" ", "-")
        print("Filename: " + print_filename)

        return document

    def enter(self):
        print("Entered!")

        # Get the selection model from the table view
        selection_model = self.tableView.selectionModel()

        if selection_model.hasSelection():

            # Get the selected rows
            selected_rows = selection_model.selectedRows()

            # Print the values of all cells in the selected rows
            for index in selected_rows:
                row = index.row()
                for column in range(self.model.columnCount()):
                    # Index of the cell in the model
                    cell_index = self.model.index(row, column)
                    # Value of the cell
                    cell_value = self.model.data(cell_index)
                    print(cell_value)

        self.closeUI()

    def closeUI(self):
        print("\nClosed!")
        self.close()

    def on_focusChanged(self, old, new):
        # If the new focus widget is not the table view, clear the selection
        if new != self.tableView and new != self.enterBtn and new != self.printBtn and old == self.tableView and old == self.enterBtn and old == self.printBtn :
            self.tableView.clearSelection()

    def prev(self):
        print("Previous!")
        self.model.clear()
        global current_page

        if current_page > 0 and current_page != 1:
            self.prevBtn.setDisabled(False)
            self.prevBtn.setStyleSheet("background-color: blue; color: white;")
            self.nextBtn.setDisabled(False)
            self.nextBtn.setStyleSheet("background-color: blue; color: white;")
            current_page = current_page - 1
            self.populate_tableview(current_page);
            self.totalLabel.setText(
                "Total Records: " + str(rows_per_page * (current_page + 1)) + "/" + str(len(self.search_isUsed())))
        elif current_page == 1:
                self.prevBtn.setDisabled(True)
                self.prevBtn.setStyleSheet("background-color: #5D8AA8; color: darkgray;")
                self.nextBtn.setDisabled(False)
                self.nextBtn.setStyleSheet("background-color: blue; color: white;")
                current_page = current_page - 1
                self.populate_tableview(current_page);

    def next(self):
            print("Next Page!")
            self.model.clear()

            global current_page
            global total_pages

            total_pages = math.ceil(len(self.search_isUsed())/rows_per_page)
            # print("Total: ", total_pages)

            if current_page < total_pages - 1:
                current_page = current_page + 1
                self.populate_tableview(current_page)

                if current_page == total_pages - 1:
                    self.totalLabel.setText("Total Records: " + str(len(self.search_isUsed())) + "/" + str(len(self.search_isUsed())))
                    self.nextBtn.setDisabled(True)
                    self.nextBtn.setStyleSheet("background-color: #5D8AA8; color: darkgray;")

            self.prevBtn.setDisabled(False)
            self.prevBtn.setStyleSheet("""
                QPushButton{
                background-color: blue;
                color: white;
                }
                QPushButton::hover{
                    background-color: #4169E1;
                    border: 2px solid white;
                    padding: 4px;
                }
            """)

    def search_isUsed(self):

        if isUsed:
            return search_result
        else:
            return result

    def search_isEmpty(self):
        global isUsed
        global current_page

        current_text = self.searchTxt.text()
        self.searchTxt.setText(current_text.upper())

        if self.searchTxt.text() == "":
            isUsed = False
            current_page - 0
            self.model.clear()
            self.nextBtn.setDisabled(False)
            self.nextBtn.setStyleSheet("background-color: blue; color: white;")
            self.populate_tableview(current_page)

    def populate_tableview(self, page):

        print("Current Page: ", current_page+1)

        global total_pages

        total_pages = math.ceil(len(self.search_isUsed()) / rows_per_page)
        print("Total Length: ", len(self.search_isUsed()))

        # Set the column headers
        headers = [column[0] for column in cursor.description]
        self.model.setHorizontalHeaderLabels(headers)

        startIndex = page * rows_per_page
        endIndex = startIndex + rows_per_page

        # Populate the model with the selected rows
        for index, row in enumerate(self.search_isUsed()):
            item_row = []
            if startIndex <= index < endIndex:
                # print("Index: ", index)
                for num, column in enumerate(row):
                    # item = QStandardItem(str(value))
                    # if column == 1:  # Specify the desired column index
                    #     item.setTextAlignment(Qt.AlignLeft)
                    # else:
                    #     item.setTextAlignment(Qt.AlignCenter)
                    # item_row.append(item)
                    # print("Column " + str(num+1) + ": " + str(column))
                    # Set the delegate for the specific column
                    item = QStandardItem(str(column))
                    item.setTextAlignment(Qt.AlignCenter)
                    item_row.append(item)

                self.model.appendRow(item_row)

            if index == len(self.search_isUsed())+1 and isUsed:
                self.nextBtn.setDisabled(True)
                self.nextBtn.setStyleSheet("background-color: #5D8AA8; color: darkgray;")
            elif current_page == total_pages - 1:
                self.nextBtn.setDisabled(True)
                self.nextBtn.setStyleSheet("background-color: #5D8AA8; color: darkgray;")

        # Set fixed size for rows and columns
        for row in range(self.model.rowCount()):
            self.tableView.setRowHeight(row, 50)  # Set row height to 50 pixels

        # Set fixed size for specific headers
        header = self.tableView.horizontalHeader()
        header.resizeSection(0, 30)
        header.setSectionResizeMode(QHeaderView.Fixed)
        header.setHighlightSections(False)
        header.setStyleSheet("""
            QHeaderView::section {
                font-size: 18px; 
                background-color: orange; 
                font-weight: bold; 
                text-align: center;
                border: 1px solid #6c6c6c;
                margin: 1px;
            }
        """)

        self.tableView.setShowGrid(False)
        self.tableView.setFont(self.font)
        self.tableView.setStyleSheet("""
            QTableView {
                border: none;
            }
        """)
        self.tableView.setColumnWidth(0, 150)
        self.tableView.setColumnWidth(1, 350)
        self.tableView.setColumnWidth(2, 160)
        self.tableView.setColumnWidth(3, 60)
        self.tableView.setColumnWidth(4, 80)
        self.tableView.setColumnWidth(5, 120)
        self.tableView.setColumnWidth(6, 80)

        if len(self.search_isUsed()) <= rows_per_page:
            self.totalLabel.setText("Total Records: " + str(len(self.search_isUsed())) + "/" + str(len(self.search_isUsed())))
        elif len(self.search_isUsed()) > rows_per_page:
            self.totalLabel.setText(
                "Total Records: " + str(rows_per_page*(current_page+1)) + "/" + str(len(self.search_isUsed())))


if __name__ == "__main__":
    app = PyQt5.QtWidgets.QApplication([])
    window = MainWindow()
    # window.setFixedSize(1040, 600)  # Set the desired width and height
    window.show()
    app.exec_()

    # cursor.close()
    conn.close()
